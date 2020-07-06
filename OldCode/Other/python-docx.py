from docx import Document
import os
f=open('GSM_eval.csv', 'r')
n=0
for i in f.readlines():
    print(i)

    n=n+1
    if n==10:
        break
    cells=i.split(',')



    document = Document()  # 首先这是包的主要接口，这应该是利用的设计模式的一种，用来创建docx文档，里面也可以包含文档路径(d:\\2.docx)

    document.add_heading('一、评估结果汇总：', 3)  # 这里是给文档添加一个标题，0表示 样式为title，1则为忽略，其他则是Heading{level},具体可以去官网查;

    p = document.add_paragraph('   针对聚类问题'+cells[1]+'，'+cells[3]+'小区进行1类详细问题进行评估： ')  # 这里是添加一个段落
#        p.add_run('bold').bold = True  # 这里是在这个段落p里文字some后面添加bold字符
#        p.add_run(' and some ')
#       p.add_run('italic.').italic = True
    table1 = document.add_table(rows=6, cols=4,style='Light List Accent 1')
    hdr_cells1 = table1.rows[0].cells
    hdr_cells1[0].text='详细问题点类型'
    hdr_cells1[1].text='评估数据粒度'
    hdr_cells1[2].text = '已解决/未解决'
    hdr_cells1[3].text = '达标项'

    hdr_cells1 = table1.rows[1].cells
    hdr_cells1[0].text='GSM高质差'
    hdr_cells1[1].text='周'
    hdr_cells1[2].text = '已解决'
    hdr_cells1[3].text = '上/下行质差话务比例至5%以下'


    document.add_heading('二、评估数据明细：', 3)
    p2=document.add_paragraph('   1.	GSM高质差问题：（数据提供人：陈俊晖） ')
    table2 = document.add_table(rows=5, cols=10, style='Light List Accent 1')
    c0=table2.rows[0].cells
    c1=table2.rows[1].cells
    c2=table2.rows[2].cells

    c3=c1[9].merge(c2[9])

    c0[0].text='评估指标名称'
    c0[1].text ='达标门限'
    c0[2].text ='2018/8/1'
    c0[3].text ='2018/8/2'
    c0[4].text ='2018/8/3'
    c0[5].text ='2018/8/4'
    c0[6].text ='2018/8/5'
    c0[7].text ='2018/8/6'
    c0[8].text ='2018/8/7'
    c0[9].text ='达标总天数'

    c1[0].text='上行质差话务比例'
    c1[1].text ='5%以下'
    c1[2].text =cells[7]
    c1[3].text =cells[8]
    c1[4].text =cells[9]
    c1[5].text =cells[10]
    c1[6].text =cells[11]
    c1[7].text =cells[12]
    c1[8].text =cells[13]
    c1[9].text =cells[14]

    c2[0].text='下行质差话务比例'
    c2[1].text ='5%以下'
    c2[2].text =cells[15]
    c2[3].text =cells[16]
    c2[4].text =cells[17]
    c2[5].text =cells[18]
    c2[6].text =cells[19]
    c2[7].text =cells[20]
    c2[8].text =cells[21]
    c2[9].text =cells[22]


    p3 = document.add_paragraph('要求：')
    p4 = document.add_paragraph('1、针对周粒度性能问题点，需提取连续7天指标，4天以上达标才能通过；')







    document.save(cells[1]+'.docx')  # 保存这